"""
Complete working example: Generate 2Penguins client proposal

This is a PRODUCTION-READY script that generates real Word documents
with company branding, professional formatting, and all necessary sections.

Usage:
    python complete_proposal_generator.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import os


class ProposalGenerator:
    """Professional proposal generator with company branding."""

    COMPANY_COLORS = {
        '2penguins': {
            'primary': RGBColor(30, 64, 175),    # Blue
            'secondary': RGBColor(245, 158, 11),  # Orange
            'text': RGBColor(31, 41, 55)          # Dark gray
        },
        'ws': {
            'primary': RGBColor(220, 38, 38),     # Red
            'secondary': RGBColor(8, 145, 178),   # Cyan
            'text': RGBColor(15, 23, 42)          # Dark blue-gray
        }
    }

    def __init__(self, company: str = '2penguins'):
        """Initialize generator with company branding."""
        self.company = company
        self.colors = self.COMPANY_COLORS[company]
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """Configure document margins and default styles."""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

    def add_title_page(self, client_name: str, project_name: str):
        """Add professional title page."""
        # Main title
        title = self.doc.add_heading(f'Proposal for {client_name}', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in title.runs:
            run.font.size = Pt(28)
            run.font.color.rgb = self.colors['primary']
            run.font.name = 'Arial'

        # Subtitle
        subtitle_p = self.doc.add_paragraph()
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_p.add_run(project_name)
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.color.rgb = self.colors['secondary']
        subtitle_run.font.name = 'Arial'
        subtitle_run.bold = True

        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Date
        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_p.add_run(f'Date: {date.today().strftime("%B %d, %Y")}')
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(107, 114, 128)

        # Company name
        self.doc.add_paragraph()
        company_p = self.doc.add_paragraph()
        company_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_p.add_run('2Penguins GmbH' if self.company == '2penguins' else 'WS Agency')
        company_run.font.size = Pt(14)
        company_run.bold = True

        # Page break
        self.doc.add_page_break()

    def add_section(self, title: str, content: str = None):
        """Add section heading with company branding."""
        heading = self.doc.add_heading(title, level=1)

        for run in heading.runs:
            run.font.size = Pt(18)
            run.font.color.rgb = self.colors['primary']
            run.font.name = 'Arial'

        if content:
            p = self.doc.add_paragraph(content)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.colors['text']

    def add_bullet_list(self, items: list):
        """Add professionally styled bullet list."""
        for item in items:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)

            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = self.colors['text']

    def add_budget_table(self, budget: float):
        """Add professional budget breakdown table."""
        # Calculate breakdown
        development = budget * 0.65
        design = budget * 0.20
        pm = budget * 0.15

        # Create table
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Item'
        header_cells[1].text = 'Amount'

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)

        # Data rows
        table.rows[1].cells[0].text = 'Development & Implementation'
        table.rows[1].cells[1].text = f'€{development:,.2f}'

        table.rows[2].cells[0].text = 'Design & UX'
        table.rows[2].cells[1].text = f'€{design:,.2f}'

        table.rows[3].cells[0].text = 'Project Management'
        table.rows[3].cells[1].text = f'€{pm:,.2f}'

        # Total row
        total_cells = table.rows[4].cells
        total_cells[0].text = 'Total Investment'
        total_cells[1].text = f'€{budget:,.2f}'

        for cell in total_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = self.colors['primary']

    def save(self, output_path: str):
        """Save document to file."""
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        return output_path


def generate_2penguins_proposal():
    """Generate complete 2Penguins client proposal."""

    # Initialize generator
    generator = ProposalGenerator(company='2penguins')

    # Client data
    client_name = "ACME Corporation"
    project_name = "Digital Signage System"
    budget = 45000.00

    deliverables = [
        "Custom web-based digital signage platform",
        "Content management system (CMS) for easy updates",
        "Responsive design for various screen sizes",
        "Real-time content synchronization across multiple displays",
        "Analytics dashboard for content performance tracking",
        "6 months of technical support and maintenance"
    ]

    timeline_items = [
        "Week 1-2: Discovery & Requirements Analysis",
        "Week 3-4: Design & Architecture",
        "Week 5-8: Development & Implementation",
        "Week 9: Testing & Quality Assurance",
        "Week 10: Deployment & Training",
        "Week 11-12: Support & Optimization"
    ]

    # Build document
    generator.add_title_page(client_name, project_name)

    generator.add_section(
        'Executive Summary',
        f'2Penguins is pleased to present this proposal for the {project_name} project. '
        f'We will deliver a comprehensive digital signage solution that enables {client_name} '
        f'to manage and display dynamic content across multiple screens efficiently.'
    )

    generator.add_section('Project Deliverables')
    generator.add_bullet_list(deliverables)

    generator.add_section('Project Timeline')
    generator.add_bullet_list(timeline_items)

    generator.add_section('Investment')
    generator.doc.add_paragraph(
        'The total investment for this project is broken down as follows:'
    )
    generator.add_budget_table(budget)

    generator.add_section(
        'Why 2Penguins?',
        'With over 10 years of experience in digital solutions, 2Penguins has successfully '
        'delivered 200+ projects for clients across Europe. Our team combines technical '
        'excellence with creative design to build solutions that drive business results.'
    )

    generator.add_section(
        'Next Steps',
        'We look forward to partnering with you on this exciting project. '
        'Please review this proposal and let us know if you have any questions. '
        'We are ready to begin work immediately upon your approval.'
    )

    # Save document
    output_path = f'proposals/acme_proposal_{date.today().isoformat()}.docx'
    final_path = generator.save(output_path)

    print(f"✅ Proposal generated successfully!")
    print(f"📄 Location: {final_path}")
    print(f"💰 Total budget: €{budget:,.2f}")
    print(f"📊 Sections: 6")
    print(f"📝 Deliverables: {len(deliverables)}")

    return final_path


if __name__ == '__main__':
    # Generate proposal
    proposal_path = generate_2penguins_proposal()

    # Verify file exists
    if os.path.exists(proposal_path):
        file_size = os.path.getsize(proposal_path) / 1024  # KB
        print(f"✅ File verified: {file_size:.1f} KB")
    else:
        print("❌ Error: File not created")
